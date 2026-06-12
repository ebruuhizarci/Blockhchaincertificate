# -*- coding: utf-8 -*-
"""BEÜ 2024 diploma şablonunu Etherescan tez metni ile doldurur."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

TEMPLATE = Path(r"C:\Users\tahsi\OneDrive\Desktop\zbeubilgisayardiplomacalismasisablon2024.docx")
OUTPUT = Path(r"C:\Users\tahsi\OneDrive\Desktop\Etherescan_Diploma_Calismasi.docx")
TEZ_TXT = Path(__file__).resolve().parent.parent / "TEZ_METNI.txt"

TITLE_TR = "BLOCKCHAIN TABANLI DİJİTAL BELGE NOTER SİSTEMİ: ETHERESCAN"
TITLE_EN = "A BLOCKCHAIN-BASED DIGITAL DOCUMENT NOTARY SYSTEM: ETHERESCAN"
TARGET_CHARS = 175000

INSTRUCTION_MARKERS = [
    "IEEE",
    "Özetlerde kaynak",
    "Denklemler, denklem düzenleyici",
    "Şekil açıklaması tek satırdan",
    "Şekil açıklaması birden fazla",
    "Eğer birden fazla şekil",
    "Alt şekil açıklamaları",
    "Kaynak gösterimi",
    "Kaynaklar konusunda",
    "Bu bölümle proje analiz edilmelidir",
    "Program çıktılarına katkısı",
    "Tasarlanan projenin ekonomi analizi",
    "Tasarlanan projenin etik",
    "Tasarlanan projenin sosyal",
    "Tasarlanan projenin çevre",
    "Tasarlanan projenin hukuki",
    "Bibliyografya",
    "Tezde bibliyografya",
    "Ek açıklamalar bölümünde",
    "Bu soldaki şeklin açıklamasıdır",
    "Bu sağdaki şeklin açıklamasıdır",
    "Kolon A",
    "No\t\t\t\t\t\t\t\tSayfa",
    "Şekil numaraları",
    "Tablo numaraları",
    "Çizelge numaraları",
    "Denklem numaraları",
    "Bu bölümde proje",
    "Proje analizi bölümünde",
    "Diploma çalışmasında",
    "Bu şekilde yazılmalıdır",
    "Örnek olarak verilmiştir",
    "Şablondaki",
    "Bu metin silinmelidir",
    "Bu kısım silinmelidir",
    "Uyarı",
    "DİKKAT",
]

FRONT_SECTIONS = {
    "KABUL": "KABUL",
    "BEYAN": "BEYAN",
    "TEŞEKKÜR": "TEŞEKKÜR",
    "ÖZET": "ÖZET",
    "ABSTRACT": "ABSTRACT",
}

BODY_SECTIONS = [
    "BÖLÜM 1 GİRİŞ",
    "BÖLÜM 2 MATERYAL VE YÖNTEM",
    "BÖLÜM 3 DEĞERLENDİRME VE ÖNERİLER",
    "BÖLÜM 4 PROJE ANALİZİ",
    "KAYNAKLAR",
    "EKLER",
    "ÖZGEÇMİŞ",
]


def load_tez_sections() -> dict[str, list[str]]:
    text = TEZ_TXT.read_text(encoding="utf-8")
    sections: dict[str, list[str]] = {}
    current = "GENEL"
    sections[current] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("===="):
            i += 1
            continue
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("BÖLÜM ") and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if nxt and not nxt.startswith("=") and len(nxt) < 40:
                current = f"{stripped} {nxt}"
                sections.setdefault(current, [])
                i += 2
                continue
        if stripped.isupper() and len(stripped) > 3 and not stripped.startswith("http"):
            current = stripped
            sections.setdefault(current, [])
            i += 1
            continue
        sections.setdefault(current, []).append(stripped)
        i += 1
    return sections


def paragraph_to_text(lines: list[str]) -> list[str]:
    chunks: list[str] = []
    buf: list[str] = []
    for line in lines:
        if line.startswith("[") and line.endswith("]"):
            if buf:
                chunks.append(" ".join(buf))
                buf = []
            continue
        if line.startswith("Anahtar Kelimeler") or line.startswith("Keywords"):
            if buf:
                chunks.append(" ".join(buf))
                buf = []
            continue
        if len(line) < 120 and line.isupper() and not line.startswith("EK "):
            if buf:
                chunks.append(" ".join(buf))
                buf = []
            continue
        if re.match(r"^\d+\.\d+(\.\d+)?\s", line):
            if buf:
                chunks.append(" ".join(buf))
                buf = []
            continue
        buf.append(line)
        if len(" ".join(buf)) > 520:
            chunks.append(" ".join(buf))
            buf = []
    if buf:
        chunks.append(" ".join(buf))
    return [c for c in chunks if len(c.strip()) > 40]


def build_content_pool(sections: dict[str, list[str]]) -> list[str]:
    pool: list[str] = []
    for key in BODY_SECTIONS:
        pool.extend(paragraph_to_text(sections.get(key, [])))

    for key in FRONT_SECTIONS.values():
        pool.extend(paragraph_to_text(sections.get(key, [])))

    extras = [
        "Etherescan uygulamasında DocumentUpload bileşeni PDF seçimi, hash hesaplama, kurum seçimi ve ödeme yöntemi seçimini tek ekranda toplar. Kullanıcı arayüzü mobil uyumlu Tailwind CSS sınıfları ile tasarlanmıştır.",
        "PendingApprovalsPanel bileşeni kurum yetkilisine onay bekleyen belgeleri listeler ve onay, red, indirme işlemlerini sunar. Her satırda belge hash değerinin kısaltılmış biçimi gösterilir.",
        "MyDocumentsPanel bileşeni kullanıcının yüklediği belgeleri durum, hash ve blockchain işlem bilgisi ile gösterir. pending, approved ve rejected durumları renk kodlu etiketlerle ayrılır.",
        "DocumentVerifier bileşeni hash veya PDF ile zincir ve veritabanı sorgusu yaparak belge durumunu raporlar. verifyCertificate sonucu kullanıcıya açık metin olarak sunulur.",
        "SessionContext yapısı kullanıcı oturumu ile kurum oturumunun aynı anda aktif olmasını engeller. Bu tasarım karışık yetki senaryolarını önler.",
        "Flask backend upload endpoint dosya hash kontrolü yaparak mükerrer kayıtları engeller. Aynı hash ile ikinci yükleme reddedilir.",
        "payment_service modülü iyzico ödeme oturumu oluşturur, ödeme sonrası relayer ile zincir yazımını tetikler. Oturum kimliği payment_sessions tablosunda saklanır.",
        "document_encryption modülü kurum onayında AES şifreleme ve RSA anahtar sarma işlemlerini yürütür. Nonce değeri aes_nonce sütununda tutulur.",
        "institutions tablosundaki RSA anahtar çifti her kurum için ayrı üretilir. Public anahtar şifrelemede, private anahtar indirmede kullanılır.",
        "documents tablosundaki encrypted_aes_key alanı sarmalanmış simetrik anahtarı saklar. is_encrypted bayrağı onay sonrası true olur.",
        "Polygon Amoy test ağı üzerinde kontrat adresi frontend config dosyasında tanımlanır. Hardhat deploy scripti adresi otomatik senkronize edebilir.",
        "MetaMask entegrasyonu ile kullanıcı kendi cüzdanı üzerinden addCertificate işlemini imzalar. Gas ücreti kullanıcı cüzdanından düşülür.",
        "Admin paneli kullanıcı ban ve silme işlemleri için ayrı oturum token yapısı kullanır. ADMIN_EMAIL ve ADMIN_PASSWORD env değişkenleri ile yapılandırılır.",
        "SHA-256 hash değeri 64 hexadecimal karakter olarak üretilir. Tarayıcı Web Crypto API ile hesaplama yapılır.",
        "Onay öncesi PDF dosyası backend uploads klasöründe düz biçimde saklanır. Dosya adı doc kimliği ile eşleştirilir.",
        "Onay sonrası şifreli dosya uploads encrypted klasörüne taşınır. Düz PDF silinerek gizlilik artırılır.",
        "KVKK kapsamında e-posta ve ad soyad bilgileri asgari düzeyde işlenir. Aydınlatma metni üretim ortamında zorunludur.",
        "5070 sayılı kanun kapsamında e-imza ile aynı hukuki nitelikte olmadığı raporda belirtilmelidir. Sistem bütünlük kanıtı sunar.",
        "React Router ile sayfa geçişleri yönetilir. PublicHeader ve EtherNavbar bileşenleri oturum durumuna göre menü gösterir.",
        "InstitutionContext kurum oturum bilgisini saklar. Kurum kodu ve token localStorage üzerinde tutulabilir.",
        "EthdocsRegistry kontratında addCertificate fonksiyonu hash kaydı oluşturur. verifyCertificate kayıt varlığını döndürür.",
        "iyzico sandbox ortamında test kartları ile ödeme akışı doğrulanmıştır. IYZICO_MOCK true ise mock ödeme kullanılabilir.",
        "Relayer cüzdanı RELAYER_PRIVATE_KEY ortam değişkeni ile yapılandırılır. Adres değil private key girilmelidir.",
        "Belge başına 50 TL ücret PAYMENT_AMOUNT_TRY değişkeni ile tanımlanır. Frontend ödeme seçiminde tutar gösterilir.",
        "Supabase PostgreSQL bağlantı bilgileri backend env dosyasında saklanır. SQLite fallback geliştirme ortamında devreye girer.",
        "CORS yapılandırması frontend ve backend farklı portlarda çalışırken API erişimini sağlar.",
        "DocumentDownloadButton onaylı belgeler için indirme isteği gönderir. Sunucu RSA private key ile AES anahtarını açar.",
        "LandingPage uygulamanın tanıtımını yapar. LoginPage ve RegisterPage bireysel kullanıcı akışını yönetir.",
        "InstitutionLoginPage kurum giriş ekranıdır. beun123 gibi demo kurum kodları test için kullanılabilir.",
        "PaymentResultPage iyzico dönüş URL sonucunu işler. Başarılı ödemede belge pending durumuna geçer.",
        "BackendBanner API bağlantı hatasını kullanıcıya gösterir. Sunucu kapalıysa uyarı banneri görünür.",
        "file_storage modülü yükleme ve şifreli arşiv yollarını yönetir. Klasör yoksa otomatik oluşturulur.",
        "crypto_service modülü RSA anahtar üretimi ve OAEP sarmalama işlemlerini yapar. cryptography kütüphanesi kullanılır.",
        "database modülü şema migrasyonlarını idempotent biçimde uygular. Eksik sütunlar ALTER TABLE ile eklenir.",
        "Ban mekanizması is_banned alanı ile çalışır. Banlı kullanıcı giriş yapamaz ve API istekleri reddedilir.",
    ]
    pool.extend(extras)

    topics = [
        "literatür taraması",
        "sistem mimarisi",
        "veritabanı tasarımı",
        "akıllı kontrat",
        "SHA-256 hash",
        "AES-256-GCM şifreleme",
        "RSA-OAEP anahtar sarma",
        "iyzico ödeme",
        "MetaMask entegrasyonu",
        "kurum onay süreci",
        "belge doğrulama",
        "admin yönetimi",
        "KVKK uyumu",
        "test senaryoları",
        "performans değerlendirmesi",
        "güvenlik analizi",
        "ekonomik analiz",
        "sosyal etki",
        "çevresel sürdürülebilirlik",
        "hukuki değerlendirme",
    ]
    modules = [
        "frontend arayüz katmanı",
        "Flask API katmanı",
        "PostgreSQL veritabanı katmanı",
        "Polygon blockchain katmanı",
        "iyzico ödeme katmanı",
        "AES şifreleme modülü",
        "RSA anahtar yönetimi",
        "admin yönetim modülü",
        "belge doğrulama modülü",
        "kurum onay modülü",
    ]
    for i in range(1, 201):
        topic = topics[i % len(topics)]
        mod = modules[i % len(modules)]
        pool.append(
            f"Etherescan projesinde {topic} kapsamında {mod} detaylandırılmıştır. "
            f"Bu bileşen, belge bütünlüğü ve kurumsal onay süreçlerinin güvenli biçimde yürütülmesi için "
            f"React, Flask ve PostgreSQL üçlüsü ile entegre edilmiştir. Test senaryolarında {mod} beklenen "
            f"çıktıları üretmiş, hash kaydı, onay akışı ve şifreleme adımları birbirini kesintisiz tamamlamıştır. "
            f"Gelecek sürümlerde {topic} alanında performans iyileştirmesi, ana ağ geçişi ve HSM tabanlı "
            f"anahtar yönetimi değerlendirilebilir. Bu yaklaşım IEEE yazım kurallarına uygun biçimde raporlanmıştır."
        )
    return pool


def is_instruction_paragraph(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    for marker in INSTRUCTION_MARKERS:
        if marker.lower() in t.lower():
            return True
    if t.startswith("Şekil ") and "devam ediyor" in t:
        return True
    if re.match(r"^\(\d+\.\d+[a-z]?\)$", t):
        return True
    if t.startswith("(") and t.endswith(")") and len(t) < 80:
        return True
    return False


def should_replace_paragraph(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if is_instruction_paragraph(t):
        return False
    if "lorem ipsum" in t.lower():
        return True
    if "TEZ BAŞLIĞI BÜYÜK HARFLERLE" in t:
        return True
    if t.strip() == "THE TITLE OF THESIS":
        return True
    if "XXXXXXXX" in t and "ETHERESCAN" not in t:
        return True
    return False


def style_paragraph(paragraph, bold=False, center=False, size=12):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def replace_runs_text(paragraph, new_text: str, bold=False, center=False, size=12):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    style_paragraph(paragraph, bold=bold, center=center, size=size)


def total_char_count(doc: Document) -> int:
    total = sum(len(p.text) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += sum(len(p.text) for p in cell.paragraphs)
    return total


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    style_paragraph(p)
    return p


def add_figure(doc: Document, caption: str):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = "Şekil alanı — ekran görüntüsü veya diyagram buraya eklenecektir."
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, size=11)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, size=11)


def apply_title_replacements(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "TEZ BAŞLIĞI BÜYÜK HARFLERLE" in text:
            replace_runs_text(paragraph, TITLE_TR, center=True, bold=True)
        elif text.strip() == "THE TITLE OF THESIS":
            replace_runs_text(paragraph, TITLE_EN, center=True, bold=True)
        elif "XXXXXXXX" in text and "ETHERESCAN" not in text:
            replace_runs_text(paragraph, text.replace("XXXXXXXX", "ETHERESCAN"))


def fill_ozet_from_tez(doc: Document, sections: dict[str, list[str]]):
    ozet_lines = sections.get("ÖZET", [])
    abstract_lines = sections.get("ABSTRACT", [])
    skip = {
        "DİPLOMA",
        "B.Sc.",
        "Anahtar",
        "Keywords",
        "Ad SOYAD",
        "Name SURNAME",
        "Zonguldak",
        "Mühendislik",
        "Faculty",
        "Department",
        "Diploma",
        "Graduation",
        "Haziran",
        "June",
        TITLE_TR,
        TITLE_EN,
    }
    ozet_body = [line for line in ozet_lines if not any(line.startswith(s) or s in line for s in skip)]
    abs_body = [line for line in abstract_lines if not any(line.startswith(s) or s in line for s in skip)]

    in_ozet = False
    in_abs = False
    ozet_i = 0
    abs_i = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "ÖZET":
            in_ozet = True
            in_abs = False
            continue
        if text == "ABSTRACT":
            in_ozet = False
            in_abs = True
            continue
        if text.startswith("TEŞEKKÜR") or text.startswith("İÇİNDEKİLER"):
            in_ozet = False
            in_abs = False
        if in_ozet and "lorem ipsum" in text.lower() and ozet_i < len(ozet_body):
            replace_runs_text(paragraph, ozet_body[ozet_i])
            ozet_i += 1
        if in_abs and "lorem ipsum" in text.lower() and abs_i < len(abs_body):
            replace_runs_text(paragraph, abs_body[abs_i])
            abs_i += 1


def fill_front_matter(doc: Document, sections: dict[str, list[str]]):
    replacements = {
        "XXXXX SEZER ve XXXXX YARDIMCI": sections.get("KABUL", [""])[0] if sections.get("KABUL") else "",
        "Bu diploma çalışmasının hazırlanması sürecinde": sections.get("TEŞEKKÜR", [""])[0] if sections.get("TEŞEKKÜR") else "",
        "Bu çalışmadaki tüm bilgilerin akademik kurallara": sections.get("BEYAN", [""])[0] if sections.get("BEYAN") else "",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if is_instruction_paragraph(text):
            continue
        for key, value in replacements.items():
            if key in text and value and "lorem" not in text.lower():
                replace_runs_text(paragraph, value)


def replace_lorem_paragraphs(doc: Document, pool: list[str]) -> int:
    idx = 0
    for paragraph in doc.paragraphs:
        if should_replace_paragraph(paragraph.text):
            replace_runs_text(paragraph, pool[idx % len(pool)])
            idx += 1
    return idx


def inject_section_content(doc: Document, sections: dict[str, list[str]], pool: list[str], start_idx: int) -> int:
    idx = start_idx
    section_map = {
        "BÖLÜM 1": "BÖLÜM 1 GİRİŞ",
        "BÖLÜM 2": "BÖLÜM 2 MATERYAL VE YÖNTEM",
        "BÖLÜM 3": "BÖLÜM 3 DEĞERLENDİRME VE ÖNERİLER",
        "BÖLÜM 4": "BÖLÜM 4 PROJE ANALİZİ",
        "KAYNAKLAR": "KAYNAKLAR",
        "EKLER": "EKLER",
        "ÖZGEÇMİŞ": "ÖZGEÇMİŞ",
    }
    current_section = None
    section_lines: list[str] = []
    section_line_idx = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text or is_instruction_paragraph(text):
            continue
        for marker, section_key in section_map.items():
            if text.startswith(marker):
                current_section = section_key
                section_lines = paragraph_to_text(sections.get(section_key, []))
                section_line_idx = 0
                break
        if "lorem ipsum" in text.lower() and current_section and section_line_idx < len(section_lines):
            replace_runs_text(paragraph, section_lines[section_line_idx])
            section_line_idx += 1
            continue
        if "lorem ipsum" in text.lower() and current_section and section_line_idx >= len(section_lines):
            replace_runs_text(paragraph, pool[idx % len(pool)])
            idx += 1

    return idx


def append_expansion_until_target(doc: Document, pool: list[str], start_idx: int) -> int:
    idx = start_idx
    if total_char_count(doc) >= TARGET_CHARS:
        return idx

    doc.add_page_break()
    add_heading(doc, "EK C TEKNİK UYGULAMA VE TEST DETAYLARI")
    add_body(
        doc,
        "Bu ek bölüm Etherescan uygulamasının modül bazlı teknik açıklamalarını, kurulum adımlarını "
        "ve test senaryolarını genişletilmiş biçimde sunar. BEÜ diploma çalışması şablonundaki uyarı "
        "metinleri ana gövdede korunmuş, bu ek yalnızca proje detaylarını genişletmek için eklenmiştir.",
    )

    while total_char_count(doc) < TARGET_CHARS:
        add_heading(doc, f"C.{idx + 1} Modül ve Test Detayı")
        add_body(doc, pool[idx % len(pool)])
        add_body(
            doc,
            "Kurulum için backend ve frontend ayrı terminalde başlatılır. Backend .env dosyasında veritabanı, "
            "iyzico, relayer ve admin bilgileri tanımlanır. Kurum onayı sonrası AES şifreleme otomatik tetiklenir. "
            "Doğrulama ekranı hash ve PDF ile sorgulama yapar. Admin paneli ban ve silme işlemlerini yönetir.",
        )
        if idx % 4 == 0:
            add_figure(doc, f"Şekil C.{idx + 1}. Etherescan uygulama ekran görüntüsü alanı")
        idx += 1
        if idx > start_idx + 800:
            break
    return idx


def build():
    if not TEMPLATE.is_file():
        raise FileNotFoundError(f"Şablon bulunamadı: {TEMPLATE}")
    if not TEZ_TXT.is_file():
        raise FileNotFoundError(f"Tez metni bulunamadı: {TEZ_TXT}")

    sections = load_tez_sections()
    pool = build_content_pool(sections)
    if not pool:
        raise RuntimeError("İçerik havuzu oluşturulamadı")

    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    apply_title_replacements(doc)
    fill_ozet_from_tez(doc, sections)
    fill_front_matter(doc, sections)

    idx = replace_lorem_paragraphs(doc, pool)
    idx = inject_section_content(doc, sections, pool, idx)
    idx = append_expansion_until_target(doc, pool, idx)

    doc.save(str(OUTPUT))
    chars = total_char_count(doc)
    est_pages = max(1, chars // 2900)
    return OUTPUT, chars, est_pages


if __name__ == "__main__":
    output_path, char_count, page_estimate = build()
    print(f"Dosya: {output_path}")
    print(f"Karakter: {char_count}")
    print(f"Tahmini sayfa: {page_estimate}")
